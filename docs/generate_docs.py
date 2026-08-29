from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor


def shade(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_run(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading_styled(doc, text, level):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x8B, 0x3A, 0x25)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                set_run(run, 10, True, (255, 255, 255))
        shade(cell, "8B3A25")
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = table.rows[r].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run(run, 10)
            if r % 2 == 0:
                shade(cell, "F7F1EC")
    doc.add_paragraph()


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("HerHelp")
    set_run(r, 28, True, (0x8B, 0x3A, 0x25))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Technical Documentation")
    set_run(r, 16, True, (0x1C, 0x24, 0x34))

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(
        "Mobile app for women and girls in Ethiopia to find verified support "
        "after abuse, violence, harassment, or other unsafe situations.\n"
        "Document version 1.0  |  Expo + Next.js + MongoDB Atlas"
    )
    set_run(r, 11, False, (0x7A, 0x5C, 0x65))

    doc.add_paragraph()

    add_heading_styled(doc, "1. Purpose of the system", 1)
    doc.add_paragraph(
        "HerHelp is a safety information app. It does not force a user to file a report. "
        "It gives step-by-step guidance, verified service contacts, emergency numbers, "
        "and plain-language legal/safety guides. Regular users only read and call. "
        "Admins log in to the same app and can add, update, and delete that content. "
        "Content lives in MongoDB Atlas so every device sees the same data."
    )

    add_heading_styled(doc, "2. Tech stack", 1)
    add_table(
        doc,
        ["Layer", "Technology", "Role in HerHelp"],
        [
            [
                "Mobile client",
                "React Native 0.86 + Expo SDK 57 + TypeScript",
                "UI, navigation, calling the API, storing the login token",
            ],
            [
                "Navigation",
                "React Navigation 7 (native stack + bottom tabs)",
                "Auth screens vs logged-in screens; Home / Find Help / Learn / Emergency tabs",
            ],
            [
                "Backend",
                "Next.js 16 (App Router) + TypeScript",
                "REST API under /api. CORS enabled so Expo (including web) can call it",
            ],
            [
                "Database",
                "MongoDB Atlas + Mongoose",
                "Users, incidents, help services, emergency numbers, learn guides",
            ],
            [
                "Auth",
                "JWT (7 days) + bcrypt (cost 12)",
                "Register/login. Passwords stored as passwordHash only",
            ],
            [
                "Token storage",
                "expo-secure-store (native) / localStorage (web)",
                "Keep the user signed in after refresh. SecureStore does not work in the browser",
            ],
            [
                "Phone calls",
                "React Native Linking tel:",
                "Emergency and Find Help buttons open the device dialer",
            ],
        ],
    )

    add_heading_styled(doc, "3. High-level architecture", 1)
    doc.add_paragraph(
        "The project is a monorepo with two apps that talk over HTTP JSON."
    )
    for line in [
        "1. The Expo app runs on a phone, emulator, or browser (Expo web).",
        "2. Screens call functions in mobile/src/api.ts.",
        "3. Those functions fetch EXPO_PUBLIC_API_URL (default http://localhost:3000).",
        "4. Next.js API routes in server/app/api/ handle the request.",
        "5. lib/mongodb.ts opens a cached Mongoose connection to Atlas.",
        "6. Models in server/models/ read or write collections.",
        "7. JSON is returned. The app updates React state (auth + content context) and re-renders.",
    ]:
        doc.add_paragraph(line, style="List Number")

    pre = doc.add_paragraph()
    r = pre.add_run(
        "Expo UI  →  fetch /api/...  →  Next.js route  →  Mongoose  →  MongoDB Atlas\n"
        "     ←  JSON (user / token / incidents / services / guides / emergencies)"
    )
    set_run(r, 10, False, (0x1C, 0x24, 0x34))

    add_heading_styled(doc, "4. Repository structure", 1)
    add_table(
        doc,
        ["Path", "What it contains"],
        [
            ["mobile/", "Expo app: screens, navigation, API client, auth and content providers"],
            ["mobile/App.tsx", "Root: AuthProvider, ContentProvider, NavigationContainer, auth vs main stacks"],
            ["mobile/Tabs.tsx", "Bottom tabs: Home, FindHelp, Learn, Emergency"],
            ["mobile/screens/", "All user-facing pages (login, home, lists, details, editors, settings)"],
            ["mobile/src/api.ts", "Typed HTTP client for every backend endpoint"],
            ["mobile/src/auth.tsx", "Logged-in user, token, isAdmin, login/logout, restore session"],
            ["mobile/src/content.tsx", "Loads GET /api/content into shared state for all screens"],
            ["mobile/src/storage.ts", "Web vs native token storage"],
            ["mobile/src/data/", "Offline fallback lists if the API is down"],
            ["server/", "Next.js API + Mongoose models + seed data"],
            ["server/app/api/", "HTTP endpoints (auth, content, CRUD)"],
            ["server/lib/", "Mongo connect, JWT helpers, CORS JSON, requireAdmin, seed"],
            ["server/models/", "User, Incident, HelpService, Emergency, Guide schemas"],
            ["server/.env.local", "MONGODB_URI, MONGODB_DB, JWT_SECRET, ADMIN_SIGNUP_CODE (not committed)"],
        ],
    )

    add_heading_styled(doc, "5. How to run the project", 1)
    doc.add_paragraph("Two terminals are required. Atlas Network Access must allow this computer’s IP (or 0.0.0.0/0 while developing).")
    doc.add_paragraph("Backend", style="Heading 3")
    doc.add_paragraph("cd server")
    doc.add_paragraph("npm run dev")
    doc.add_paragraph("API base URL: http://localhost:3000")
    doc.add_paragraph("Mobile", style="Heading 3")
    doc.add_paragraph("cd mobile")
    doc.add_paragraph("npx expo start")
    doc.add_paragraph(
        "Press w for web, a for Android, or scan the QR code with Expo Go. "
        "On a physical phone, set mobile/.env EXPO_PUBLIC_API_URL to your PC LAN IP, "
        "for example http://192.168.1.10:3000. Android emulator often uses http://10.0.2.2:3000."
    )

    add_heading_styled(doc, "6. Application startup and global state", 1)
    doc.add_paragraph(
        "App.tsx wraps the tree in GestureHandlerRootView, AuthProvider, ContentProvider, "
        "and NavigationContainer."
    )
    doc.add_paragraph("AuthProvider (mobile/src/auth.tsx)", style="Heading 3")
    for line in [
        "On launch, it reads the saved token.",
        "If a token exists, it calls GET /api/auth/me. The user object (including role) comes from the database, not only from the JWT payload.",
        "If the token is missing or invalid, the user is treated as logged out.",
        "RootNavigator then shows AuthNavigator (login family) or MainNavigator (tabs + extra screens).",
        "isAdmin is true when user.role === \"admin\".",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_paragraph("ContentProvider (mobile/src/content.tsx)", style="Heading 3")
    for line in [
        "Calls GET /api/content and stores incidents, services, emergencies, and guides.",
        "If Atlas is unreachable, it falls back to the static lists in mobile/src/data/ so the UI still opens.",
        "refresh() is called after admin save/delete so every screen sees new data.",
        "It also calls refreshUser() so a role change in Atlas can appear after reload.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    add_heading_styled(doc, "7. Navigation map — how every page is connected", 1)
    doc.add_paragraph(
        "There are three navigation trees. The user never sees all of them at once. "
        "If there is no session, only the auth stack exists. After login, the main stack exists, "
        "and its first screen is the tab navigator named Tabs."
    )

    add_heading_styled(doc, "7.1 Auth stack (logged out)", 2)
    add_table(
        doc,
        ["Screen", "File", "Opens when", "Goes to"],
        [
            [
                "Login",
                "LoginScreen.tsx",
                "App start if not signed in",
                "Register; Forgot password; after success, MainNavigator (Home tab)",
            ],
            [
                "Register",
                "RegisterScreen.tsx",
                "Login → Register",
                "Login; after success, MainNavigator (Home tab)",
            ],
            [
                "Forgot password",
                "ForgotPasswordScreen.tsx",
                "Login → Forgot password?",
                "Back to Login (reset is not implemented yet)",
            ],
        ],
    )
    doc.add_paragraph(
        "WelcomeScreen.tsx still exists in the folder but is not registered in App.tsx. "
        "The live first screen is Login."
    )

    add_heading_styled(doc, "7.2 Bottom tabs (always visible after login)", 2)
    doc.add_paragraph(
        "Tabs.tsx is the hub. Switching tabs does not unmount the stack screens above Tabs; "
        "it only changes which tab is focused. From any tab, the top initials open Settings "
        "on the parent stack."
    )
    add_table(
        doc,
        ["Tab", "Screen file", "User actions", "Where those actions go"],
        [
            [
                "HOME",
                "HomeScreen.tsx",
                "Immediate danger banner; 4 featured incidents; See all incident types; Find Help card; Learn card",
                "tel: police number; IncidentDetail; IncidentList; FindHelp tab; Learn tab",
            ],
            [
                "FIND HELP",
                "FindHelpScreen.tsx",
                "Search, category chips, Call on a service card",
                "tel: service phone. Admin: EditService",
            ],
            [
                "LEARN",
                "LearnScreen.tsx",
                "Category chips, tap a guide card",
                "GuideDetail. Admin: EditGuide",
            ],
            [
                "EMERGENCY",
                "EmergencyScreen.tsx",
                "Tap a terracotta card / number pill",
                "tel: that number. Admin: EditEmergency",
            ],
        ],
    )

    add_heading_styled(doc, "7.3 Main stack screens (pushed on top of tabs)", 2)
    add_table(
        doc,
        ["Screen", "Params", "Opened from", "Back / next"],
        [
            [
                "IncidentList (Something happened)",
                "none",
                "Home → See all incident types",
                "← HOME returns to Home tab. Tap a row → IncidentDetail",
            ],
            [
                "IncidentDetail",
                "incidentId",
                "Home featured tile or IncidentList row",
                "← ALL INCIDENT TYPES → IncidentList. CALL → tel:. Find Help text → FindHelp tab. Admin Edit/Delete",
            ],
            [
                "GuideDetail",
                "guideId",
                "Learn list card",
                "← ALL GUIDES → Learn tab. Admin Edit/Delete",
            ],
            [
                "Settings",
                "none",
                "Top-right avatar (initials) on any chrome header",
                "← BACK. Log out clears token and returns to Login",
            ],
            [
                "EditIncident",
                "incidentId? (empty = create)",
                "Admin + Add/Edit on Home, IncidentList, IncidentDetail",
                "Save → POST/PUT /api/incidents then refresh content and goBack",
            ],
            [
                "EditService",
                "serviceId?",
                "Admin on Find Help",
                "Save → POST/PUT /api/services",
            ],
            [
                "EditEmergency",
                "emergencyId?",
                "Admin on Emergency",
                "Save → POST/PUT /api/emergencies",
            ],
            [
                "EditGuide",
                "guideId?",
                "Admin on Learn or GuideDetail",
                "Save → POST/PUT /api/guides",
            ],
        ],
    )

    add_heading_styled(doc, "7.4 Page-by-page behavior", 2)

    doc.add_paragraph("Login", style="Heading 3")
    doc.add_paragraph(
        "Terracotta brand header plus a white sheet. Fields: email, password. "
        "Submit calls loginAccount → POST /api/auth/login. On success setSession stores the JWT "
        "and user. Because user is now non-null, RootNavigator unmounts the auth stack and mounts "
        "MainNavigator, which shows the Home tab. Register and Forgot password are stack pushes."
    )

    doc.add_paragraph("Register", style="Heading 3")
    doc.add_paragraph(
        "Same visual family as Login. Fields: full name, email, phone, password (minimum 8 characters). "
        "Submit calls registerAccount → POST /api/auth/register with role user. "
        "The API creates the MongoDB user, hashes the password, returns a token, and the app "
        "logs the person in immediately (same as login). Duplicate phone or email returns 409."
    )

    doc.add_paragraph("Home", style="Heading 3")
    doc.add_paragraph(
        "Uses useAuth for the first name greeting and useContent for lists. "
        "The danger banner uses the first emergency whose name contains “police”, otherwise the first "
        "emergency, otherwise 991. Featured incidents are incidents with featured === true (max four). "
        "Each tile navigates to IncidentDetail with that MongoDB id. "
        "See all incident types pushes IncidentList. The Find Help and Learn cards call "
        "navigation.navigate on sibling tabs. Admins also see + Add incident and Edit/Delete on tiles."
    )

    doc.add_paragraph("Something happened (IncidentList)", style="Heading 3")
    doc.add_paragraph(
        "Scrolls every incident from ContentProvider. A row opens IncidentDetail. "
        "Admins can add a new incident or edit/delete existing ones. Delete calls DELETE /api/incidents/:id "
        "then refresh()."
    )

    doc.add_paragraph("Incident detail", style="Heading 3")
    doc.add_paragraph(
        "Looks up incidentId in the loaded incidents array. Shows summary, a call banner, and numbered "
        "action-plan steps. If step text contains “Find Help”, that phrase is a link to the FindHelp tab. "
        "Admins can edit the whole incident (title, summary, featured flag, steps) or delete it and return to the list."
    )

    doc.add_paragraph("Find Help", style="Heading 3")
    doc.add_paragraph(
        "Directory of HelpService documents. Search matches name, category, area, and description. "
        "Chips are All plus unique categories from live data. Call uses tel: on the service phone. "
        "Admins add/edit/delete services (name, category, description, area, hours, phone, verified)."
    )

    doc.add_paragraph("Emergency", style="Heading 3")
    doc.add_paragraph(
        "List of Emergency documents. Tapping a card opens the dialer. Admins manage name, description, and number. "
        "The Home danger banner reads from this same list, so editing Police here changes Home as well after refresh."
    )

    doc.add_paragraph("Learn", style="Heading 3")
    doc.add_paragraph(
        "List of Guide documents with category chips. Tapping a card opens GuideDetail. "
        "Admins add/edit/delete guides."
    )

    doc.add_paragraph("Guide detail", style="Heading 3")
    doc.add_paragraph(
        "Shows intro, optional heading, bullet list, and footer. Back goes to the Learn tab. "
        "Admin edit opens EditGuide with the same guideId."
    )

    doc.add_paragraph("Settings", style="Heading 3")
    doc.add_paragraph(
        "Shows initials, full name, role (User or Admin), email, and phone. Log out deletes the stored token "
        "and clears React auth state, which switches navigation back to Login."
    )

    add_heading_styled(doc, "8. Backend in detail", 1)
    doc.add_paragraph(
        "Next.js runs only as an API (plus a simple status page at /). "
        "Every route sets runtime = nodejs because Mongoose cannot run on the Edge runtime. "
        "lib/http.ts adds CORS headers so Expo web on another origin can call the API. "
        "OPTIONS handlers return 204 for preflight."
    )

    add_heading_styled(doc, "8.1 Connection and security helpers", 2)
    add_table(
        doc,
        ["Module", "Responsibility"],
        [
            [
                "lib/mongodb.ts",
                "Cached mongoose.connect. dbName from MONGODB_DB (default HERHELP). "
                "IPv4 family: 4 (helps Windows). Failed connects clear the cache so the next request retries.",
            ],
            [
                "lib/auth.ts",
                "hashPassword / verifyPassword (bcrypt). signAuthToken / verifyAuthToken (JWT, 7 days). "
                "Bearer header parsing. publicUser strips passwordHash.",
            ],
            [
                "lib/admin.ts",
                "requireAdmin: decode JWT, load User by id from Atlas, allow only role admin. "
                "Role is checked in the database, not only inside the token.",
            ],
            [
                "lib/content.ts",
                "seedContentIfEmpty inserts the original lists once if a collection has zero documents. "
                "loadAllContent returns all four lists sorted by sortOrder. publicDoc maps _id to id.",
            ],
        ],
    )

    add_heading_styled(doc, "8.2 Environment variables (server/.env.local)", 2)
    add_table(
        doc,
        ["Variable", "Meaning"],
        [
            ["MONGODB_URI", "Atlas SRV connection string"],
            ["MONGODB_DB", "Database name, typically HERHELP"],
            ["JWT_SECRET", "Secret used to sign and verify tokens"],
            ["ADMIN_SIGNUP_CODE", "Required to register extra admins after the first admin exists (register API)"],
        ],
    )
    doc.add_paragraph(
        "Never commit .env.local. Mobile uses EXPO_PUBLIC_API_URL only."
    )

    add_heading_styled(doc, "8.3 REST API catalog", 2)
    add_table(
        doc,
        ["Method", "Path", "Who", "What it does"],
        [
            ["GET", "/api/health", "Anyone", "Health check { ok, service }"],
            ["POST", "/api/auth/register", "Anyone", "Create user. Body: fullName, phone, email, password, optional role/adminCode. Returns token + user"],
            ["POST", "/api/auth/login", "Anyone", "Body: email (or phone) + password. Returns token + user"],
            ["GET", "/api/auth/me", "Bearer token", "Current user from DB (fresh role)"],
            ["GET", "/api/content", "Anyone", "Seed if empty; return incidents, services, emergencies, guides"],
            ["POST", "/api/incidents", "Admin", "Create incident (title, summary, featured, steps)"],
            ["PUT", "/api/incidents/:id", "Admin", "Replace incident"],
            ["DELETE", "/api/incidents/:id", "Admin", "Delete incident"],
            ["POST", "/api/services", "Admin", "Create help service"],
            ["PUT", "/api/services/:id", "Admin", "Update help service"],
            ["DELETE", "/api/services/:id", "Admin", "Delete help service"],
            ["POST", "/api/emergencies", "Admin", "Create emergency number"],
            ["PUT", "/api/emergencies/:id", "Admin", "Update emergency number"],
            ["DELETE", "/api/emergencies/:id", "Admin", "Delete emergency number"],
            ["POST", "/api/guides", "Admin", "Create learn guide"],
            ["PUT", "/api/guides/:id", "Admin", "Update learn guide"],
            ["DELETE", "/api/guides/:id", "Admin", "Delete learn guide"],
        ],
    )

    add_heading_styled(doc, "8.4 Auth API details", 2)
    doc.add_paragraph(
        "Register requires fullName, phone, email, and password. Email must contain @. "
        "Password length at least 8. Phone and email must be unique. "
        "role defaults to user. If role is admin: the first admin in the database may register freely; "
        "later admins must send adminCode matching ADMIN_SIGNUP_CODE. "
        "The intended production path is: register as user, then set role to admin in Atlas."
    )
    doc.add_paragraph(
        "Login finds the user by email (or phone if email is omitted). bcrypt.compare checks the password. "
        "Invalid credentials return 401. Success signs JWT { userId, role }."
    )
    doc.add_paragraph(
        "Me requires a valid Bearer token. It loads the User document so Settings and isAdmin "
        "reflect Atlas even if the JWT was issued when the role was still user — after a refresh/login."
    )

    add_heading_styled(doc, "8.5 Content API and seeding", 2)
    doc.add_paragraph(
        "GET /api/content is the single read used by the app. If Incidents, HelpServices, Emergencies, "
        "or Guides have zero documents, seed data from server/data/ (copied from the original mobile lists) "
        "is inserted once. After that, only admin CRUD changes the data. "
        "Users always read Atlas (or the mobile fallback if the request fails)."
    )

    add_heading_styled(doc, "9. Database collections", 1)
    doc.add_paragraph(
        "Atlas database name: HERHELP (or MONGODB_DB). Mongoose collection names are the pluralized model names."
    )
    add_table(
        doc,
        ["Model / collection", "Main fields", "Used by screens"],
        [
            [
                "User / users",
                "fullName, phone (unique), email (unique sparse), passwordHash, role user|admin, timestamps",
                "Login, Register, Settings, requireAdmin",
            ],
            [
                "Incident / incidents",
                "title, summary, featured, steps[{title, body}], sortOrder",
                "Home featured grid, IncidentList, IncidentDetail, EditIncident",
            ],
            [
                "HelpService / helpservices",
                "name, category, description, area, hours, phone, verified, sortOrder",
                "FindHelp, EditService",
            ],
            [
                "Emergency / emergencies",
                "name, description, number, sortOrder",
                "Emergency tab, Home danger banner, IncidentDetail call, EditEmergency",
            ],
            [
                "Guide / guides",
                "category, title, summary, intro, heading, bullets[], footer, sortOrder",
                "Learn, GuideDetail, EditGuide",
            ],
        ],
    )

    add_heading_styled(doc, "10. User vs admin", 1)
    doc.add_paragraph(
        "Both roles use the same navigation and the same GET /api/content. "
        "The difference is UI and write APIs."
    )
    for line in [
        "User: read-only. No Add/Edit/Delete buttons. Can call numbers and open details.",
        "Admin: Chrome shows ADMIN above the avatar. AdminAddButton and AdminActions appear on content screens.",
        "To promote someone: in Atlas open users, set role to \"admin\", then log out and log in (or reload so /me runs).",
        "Write routes return 401 without a token and 403 if the DB role is not admin. The app cannot bypass this by hiding buttons only.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    add_heading_styled(doc, "11. End-to-end example flows", 1)
    doc.add_paragraph("A. New survivor registers and reads guidance", style="Heading 3")
    for line in [
        "Open app → Login → Register → POST /api/auth/register → token stored.",
        "MainNavigator + ContentProvider GET /api/content.",
        "Home → Physical assault tile → IncidentDetail (incidentId).",
        "Follow action plan; tap Call; or Find Help text → FindHelp tab → Call a shelter.",
        "Avatar → Settings → Log out → token deleted → Login.",
    ]:
        doc.add_paragraph(line, style="List Number")

    doc.add_paragraph("B. Admin updates a Learn guide", style="Heading 3")
    for line in [
        "Log in as an account whose Atlas role is admin.",
        "Learn tab → Edit on a card → EditGuide.",
        "Change title/body → Save → PUT /api/guides/:id with Bearer token.",
        "requireAdmin loads User, confirms admin, Mongoose updates the document.",
        "refresh() GET /api/content → Learn and GuideDetail show the new text for every user.",
    ]:
        doc.add_paragraph(line, style="List Number")

    add_heading_styled(doc, "12. Important operational notes", 1)
    for line in [
        "Atlas Network Access must include this IP or 0.0.0.0/0. Otherwise login/register/content return 500 with MongooseServerSelectionError.",
        "If the Atlas user password contains special characters, URL-encode them in MONGODB_URI.",
        "Forgot password is UI-only; there is no reset email API yet.",
        "JWT lasts 7 days. Changing JWT_SECRET invalidates all sessions.",
        "First GET /api/content after an empty database seeds default Ethiopian resources; deleting all documents and calling GET again would re-seed.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    add_heading_styled(doc, "13. Summary", 1)
    doc.add_paragraph(
        "HerHelp is an Expo client and a Next.js JSON API sharing one MongoDB Atlas database. "
        "Pages are connected through React Navigation: an auth stack before login, four tabs after login, "
        "and stack screens for lists, details, settings, and admin editors. "
        "Reads go through GET /api/content. Writes are admin-only CRUD routes. "
        "That is the full path from a tap on the phone to a document in Atlas and back."
    )

    footer = doc.add_paragraph()
    r = footer.add_run("End of document.")
    set_run(r, 10, False, (0x7A, 0x5C, 0x65))

    out = Path(__file__).resolve().parent / "HerHelp-Technical-Documentation.docx"
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
